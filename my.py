from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Performance Evaluation System for ASTU University Employees', 0)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph("The Performance Evaluation System aims to streamline the task assignment and employee "
                  "evaluation process at ASTU University. It addresses the challenges of manually managing tasks, "
                  "tracking progress, and conducting employee evaluations through a digital platform. The system "
                  "automates evaluations, task assignments, and reporting, providing a more efficient solution for "
                  "employees, Unit Leaders, and Work Unit Managers.")

# Stakeholders
doc.add_heading('2. Stakeholders', level=1)
doc.add_paragraph("1. Work Unit Manager: Manages multiple units and assigns half-year plans to the units.\n"
                  "2. Unit Leader: Oversees the employees within a unit, assigns tasks, conducts evaluations, and "
                  "generates reports.\n"
                  "3. Employees: Execute assigned tasks and undergo performance evaluations.")

# System Objectives
doc.add_heading('3. System Objectives', level=1)

doc.add_heading('3.1 General Goal', level=2)
doc.add_paragraph("The system's general goal is to automate task assignment, track employee progress, and "
                  "digitize the performance evaluation process.")

doc.add_heading('3.2 Specific Objectives', level=2)
doc.add_paragraph("1. Automate task assignment and tracking to improve accountability.\n"
                  "2. Enable Unit Leaders to create and update half-year plans.\n"
                  "3. Automate the three-part performance evaluation process, including self, peer, and leader evaluations.\n"
                  "4. Facilitate report generation for employee performance by Unit Leaders.\n"
                  "5. Ensure secure role-based login and data access.")

# Requirements
doc.add_heading('4. Requirements', level=1)

doc.add_heading('4.1 Functional Requirements', level=2)
doc.add_paragraph("1. User Registration and Login: Unit Leaders register employees. Login is role-based with email and password.\n"
                  "2. Task Assignment: Unit Leaders assign tasks with details such as title, description, deadline, and priority.\n"
                  "3. Half-Year Plan Creation: Work Unit Managers assign half-year plans to units.\n"
                  "4. Performance Evaluation: Conduct three-part evaluations by employees, peers, and Unit Leaders.\n"
                  "5. Report Generation: Unit Leaders generate reports based on employee performance and evaluations.")

doc.add_heading('4.2 Non-Functional Requirements', level=2)
doc.add_paragraph("1. Performance: The system should handle up to 500 simultaneous users with response times under 2 seconds.\n"
                  "2. Scalability: The system should scale to accommodate more users and roles without downtime.\n"
                  "3. Security: Role-based access control and data encryption to protect sensitive information.\n"
                  "4. Usability: Intuitive UI with a 3-click rule to access primary functions.\n"
                  "5. Reliability: 99.9% uptime and daily backups with recovery times under 1 hour.")

# Design Goals and Architecture
doc.add_heading('5. Design Goals and Architecture', level=1)

doc.add_heading('5.1 Design Goals', level=2)
doc.add_paragraph("1. User-friendliness\n"
                  "2. Scalability\n"
                  "3. Security\n"
                  "4. Performance\n"
                  "5. Reliability\n"
                  "6. Automation")

doc.add_heading('5.2 System Architecture', level=2)
doc.add_paragraph("1. Presentation Layer: Web and mobile interfaces for users (Employees, Unit Leaders, Work Unit Managers).\n"
                  "2. Application Layer: Handles task management, evaluation, and report generation.\n"
                  "3. Data Layer: Stores user data, tasks, plans, and evaluations.\n"
                  "4. Security Layer: Provides secure authentication and authorization.\n"
                  "5. Integration Layer: Enables communication between the frontend and backend systems.")

# Use Cases
doc.add_heading('6. Use Cases', level=1)

doc.add_paragraph("1. Task Assignment: Unit Leader assigns tasks to employees with details like title, description, and deadline.\n"
                  "2. Half-Year Plan Assignment: Work Unit Manager assigns half-year plans to units.\n"
                  "3. Performance Evaluation: Employees complete self-evaluations, peers provide feedback, and Unit Leaders evaluate tasks.\n"
                  "4. Report Generation: Unit Leader generates reports on employee performance based on evaluations and task completion.")

# Database Design
doc.add_heading('7. Database Design', level=1)

doc.add_paragraph("1. Users (Employees, Unit Leaders, Work Unit Managers): Stores user information like name, role, email, and unit.\n"
                  "2. Units: Contains details of units managed by Unit Leaders.\n"
                  "3. Tasks: Stores task data such as title, description, deadlines, and assigned employees.\n"
                  "4. Half-Year Plans: Stores plans assigned by Work Unit Managers to units.\n"
                  "5. Evaluations: Stores self, peer, and Unit Leader evaluations for employees.\n"
                  "6. Reports: Contains generated performance reports for employees.")

# Deployment
doc.add_heading('8. Deployment', level=1)

doc.add_paragraph("1. Frontend: Deployed on cloud hosting services (e.g., AWS, Azure).\n"
                  "2. Backend: Hosted on cloud infrastructure with APIs handling business logic and data processing.\n"
                  "3. Database: Cloud-based database (e.g., PostgreSQL, MySQL) for storing performance data.\n"
                  "4. Security: SSL/TLS encryption and role-based access control for secure data handling.\n"
                  "5. Scalability: Auto-scaling for increased usage during evaluation periods.")

# Save the document
file_path = "Performance_Evaluation_System_ASTU.docx"
doc.save(file_path)

file_path
